from pathlib import Path
from datetime import datetime

from docx import Document


def export_report(
    state: dict,
    output_dir: str = "reports"
) -> str:
    """
    Export Research Summary to DOCX.
    """

    summary = state["research_summary"]

    query = state["query"]

    research_plan = (
        state["research_plan"]
    )

    information_gaps = (
        state["information_gaps"]
    )

    critic_feedback = (
        state["critic_feedback"]
    )

    critic_score = (
        state["critic_score"]
    )

    memory_report = (
        state["memory_report"]
    )

    memory_recommendations = (
        state["memory_recommendations"]
    )

    Path(
        output_dir
    ).mkdir(
        parents=True,
        exist_ok=True
    )

    timestamp = (
        datetime.now()
        .strftime(
            "%Y%m%d_%H%M%S"
        )
    )

    output_path = (
        Path(output_dir)
        /
        f"research_report_{timestamp}.docx"
    )

    document = Document()

    # ------------------------------------------
    # Title
    # ------------------------------------------

    document.add_heading(
        "D.A.R.I.A. Research Intelligence Report",
        level=0
    )

    document.add_paragraph(
        f"Generated On: "
        f"{datetime.now():%Y-%m-%d %H:%M:%S}"
    )

    # ------------------------------------------
    # Query
    # ------------------------------------------

    document.add_heading(
        "User Query",
        level=1
    )

    document.add_paragraph(
        query
    )

    # ------------------------------------------
    # Executive Summary
    # ------------------------------------------

    document.add_heading(
        "Executive Summary",
        level=1
    )

    document.add_paragraph(
        summary[
            "executive_summary"
        ]
    )

    # ------------------------------------------
    # Key Findings
    # ------------------------------------------

    document.add_heading(
        "Key Findings",
        level=1
    )

    for item in summary[
        "key_findings"
    ]:

        document.add_paragraph(
            item,
            style="List Bullet"
        )

    # ------------------------------------------
    # Detailed Analysis
    # ------------------------------------------

    document.add_heading(
        "Detailed Analysis",
        level=1
    )

    document.add_paragraph(
        summary[
            "detailed_analysis"
        ]
    )

    # ------------------------------------------
    # Sources
    # ------------------------------------------

    document.add_heading(
        "Sources",
        level=1
    )

    for source in summary[
        "source_summary"
    ]:

        document.add_paragraph(
            source,
            style="List Bullet"
        )

    # ------------------------------------------
    # Confidence
    # ------------------------------------------

    document.add_heading(
        "Confidence Assessment",
        level=1
    )

    document.add_paragraph(
        summary[
            "confidence_assessment"
        ]
    )

    # ------------------------------------------
    # Appendix A — Research Plan
    # ------------------------------------------

    document.add_page_break()

    document.add_heading(
        "Appendix A - Research Plan",
        level=1
    )

    document.add_paragraph(
        f"Route Selected: "
        f"{state['route']}"
    )

    for step in research_plan:

        document.add_paragraph(
            step,
            style="List Bullet"
        )

    # ------------------------------------------
    # Appendix B — Information Gaps
    # ------------------------------------------

    document.add_heading(
        "Appendix B - Information Gaps",
        level=1
    )

    for gap in information_gaps:

        document.add_paragraph(
            gap,
            style="List Bullet"
        )

    # ------------------------------------------
    # Appendix C — Critic Evaluation  
    # ------------------------------------------

    document.add_heading(
        "Appendix C - Critic Evaluation",
        level=1
    )

    document.add_paragraph(
        f"Critic Score: "
        f"{critic_score}"
    )

    document.add_paragraph(
        critic_feedback
    )

    # ------------------------------------------
    # Appendix D — Memory Insights  
    # ------------------------------------------

    if state["memory_hit"]:
        document.add_heading(
            "Appendix D - Memory Insights",
            level=1
        )

        document.add_paragraph(
            memory_report
        )

        for item in memory_recommendations:

            document.add_paragraph(
                item,
                style="List Bullet"
            )

    # ------------------------------------------
    # Appendix E — Research Metadata  
    # ------------------------------------------

    document.add_heading(
        "Appendix E - Research Metadata",
        level=1
    )

    document.add_paragraph(
        f"Iterations: "
        f"{state['iteration_count']}"
    )

    document.add_paragraph(
        f"Memory Hit: "
        f"{state['memory_hit']}"
    )

    document.add_paragraph(
        f"Memory Distance: "
        f"{state['memory_distance']:.3f}"
    )

    document.save(
        output_path
    )

    return str(
        output_path
    )